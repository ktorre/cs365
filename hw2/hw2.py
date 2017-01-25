#   Author : Kevin De La Torre

from sys import argv
from docx import Document
import sys



def main():
    if ( len( argv ) > 2 ):
        instructionDict = loadInstructionDictionary()
        registerDict = loadRegisterDictionary()
        sys.stdout.write( "Processing..." )
        document = Document()
        writeCoverPage( document )
        pcFile = open( argv[ 1 ], 'r' )
        irFile = open( argv[ 2 ], 'r' )
        
        for pc in pcFile:
            irLine = irFile.readline()
            attributes = decode( irLine, instructionDict, registerDict, pc )
            #convertInstruction( irLine, instructionDict, registerDict, pc )
            createTable( document, attributes )
            document.add_page_break()

        print( "Done." )
        document.save( "CS365_Homework2.docx" )
    else:
        print( "Usage: {0} <pc> <ir>".format( argv[ 0 ] ) )

'''
def printDictionary( dict ):
    for key in dict.keys():
        print( "Key: {0}".format( key ) )
        print( "Value(s): {0}".format( dict.get( key ) ) )
'''

def loadRegisterDictionary():
    sys.stdout.write( "Loading registers..." )
    file = open( "registers.dat", 'r' )
    registerDict = {}
    for line in file:
        line = line.split( ',', 2 )
        registerDict[ line[ 0 ] ] = line[ 1 ]
    print( "Done." )
    return registerDict

def loadInstructionDictionary():
    sys.stdout.write( "Loading instructions..." )
    file = open( "instructions.dat", 'r' )
    instructionDict = {}
    for line in file:
        line = line.split( ',', 4 )
        instructionDict[ line[ 0 ] ] = line[ 1:4 ]
    print( "Done." )

    return instructionDict


def createTable( document, attributes ):
    table = document.add_table( rows = 7, cols = 2 )
    # Fill in left side of table
    table.rows[ 0 ].cells[ 0 ].text = "PC ( Program Counter )"
    table.rows[ 1 ].cells[ 0 ].text = "Instruction Register"
    table.rows[ 2 ].cells[ 0 ].text = "PC ( Post-IR retrieval )"
    table.rows[ 3 ].cells[ 0 ].text = "Instruction Type"
    table.rows[ 4 ].cells[ 0 ].text = "Addressing Mode"
    table.rows[ 5 ].cells[ 0 ].text = "MIPS Instruction"
    table.rows[ 6 ].cells[ 0 ].text = "PC to be executed next"

    # Fill in right side of table, had to do these separate cause different lines required different formatting
    table.rows[ 0 ].cells[ 1 ].text = "0x{:06X}".format( attributes[ 0 ] )
    table.rows[ 1 ].cells[ 1 ].text = "0x{:08X}".format( attributes[ 1 ] )
    table.rows[ 2 ].cells[ 1 ].text = "0x{:06X}".format( attributes[ 2 ] )
    table.rows[ 3 ].cells[ 1 ].text = attributes[ 3 ]
    table.rows[ 4 ].cells[ 1 ].text = attributes[ 4 ]
    table.rows[ 5 ].cells[ 1 ].text = attributes[ 5 ]
    table.rows[ 6 ].cells[ 1 ].text = "0x{:06X}".format( attributes[ 6 ] )

#def convertInstruction( currentInstruction ):
    

def decode( instruction, instructions, registers, pc ):
    pc = int( pc, 16 )

    instructionName = instruction.split( ' ' )[ 0 ]
    instruction
    instructionValues = instructions.get( instructionName )
    print( "Name: ", instructionName )
    print( "Values: ", instructionValues )
    print()
    
    ir = 0x2012001E
    pcPost = 0x4000004
    type = "I-Type"
    addressMode = "Immediate Addressing"
    mipsInstruction = "addi $s2, $zero, 30"
    pcToExec = 0x4000004
    hex = 0x00000000
    return ( pc, ir, pcPost, type, addressMode, mipsInstruction, pcToExec, hex )


def writeCoverPage( document ):
    document.add_heading( "CS365 - Homework 2" )
    document.add_paragraph( "Kevin De La Torre ( 010806250 )" )
    document.add_page_break()

if __name__ == "__main__":
    main()
